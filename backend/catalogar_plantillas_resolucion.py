"""Cataloga modelos Word de Secretaría sin modificar los archivos originales."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import unicodedata
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document


ROOT_PREDETERMINADO = Path("/opt/CARPETA DE SECRETARÍA ACADEMICA")
SALIDA_PREDETERMINADA = Path("/opt/sistema_posgrado/data/plantillas_resolucion")

NOMBRES_PASO = {
    1: "Nombramiento de asesor",
    2: "Dictamen de proyecto de tesis",
    3: "Inscripción del proyecto de tesis",
    4: "Declarado apto al grado",
    5: "Dictamen de tesis",
    6: "Sustentación de tesis",
    7: "Elevación al VRAC / diploma",
}


def texto_normalizado(valor: str | None) -> str:
    base = unicodedata.normalize("NFKD", str(valor or ""))
    ascii_texto = "".join(caracter for caracter in base if not unicodedata.combining(caracter))
    return re.sub(r"\s+", " ", ascii_texto.upper()).strip()


def _texto_documento(documento: Document) -> str:
    partes = [parrafo.text for parrafo in documento.paragraphs]
    partes.extend(
        celda.text
        for tabla in documento.tables
        for fila in tabla.rows
        for celda in fila.cells
    )
    for seccion in documento.sections:
        partes.extend(parrafo.text for parrafo in seccion.header.paragraphs)
        partes.extend(parrafo.text for parrafo in seccion.footer.paragraphs)
    return "\n".join(parte for parte in partes if parte and parte.strip())


def _huella_formato(documento: Document) -> str:
    """Agrupa copias con la misma construcción visual aunque cambien sus datos."""
    elementos = []
    for parrafo in documento.paragraphs:
        elementos.append(
            (
                "P",
                parrafo.style.style_id if parrafo.style else None,
                str(parrafo.alignment),
                tuple(
                    (
                        bool(run.bold),
                        bool(run.italic),
                        bool(run.underline),
                        round(run.font.size.pt, 1) if run.font.size else None,
                        run.font.name,
                    )
                    for run in parrafo.runs
                ),
            )
        )
    for tabla in documento.tables:
        elementos.append(("T", len(tabla.rows), len(tabla.columns), tabla.style.style_id if tabla.style else None))
    for seccion in documento.sections:
        elementos.append(
            (
                "S",
                seccion.page_width,
                seccion.page_height,
                seccion.top_margin,
                seccion.bottom_margin,
                seccion.left_margin,
                seccion.right_margin,
            )
        )
    return hashlib.sha256(repr(elementos).encode("utf-8")).hexdigest()[:16]


PATRONES_PASO = {
    1: ("NOMBRAMIENTO DE ASESOR", "NOMBRAR COMO ASESOR", "DESIGNAR COMO ASESOR"),
    2: ("DICTAMEN DE PROYECTO", "DICTAMINANTE DEL PROYECTO", "DICTAMINANTES DEL PROYECTO"),
    3: ("INSCRIPCION DEL PROYECTO", "INSCRIPCION DE PROYECTO", "INSCRIBIR EL PROYECTO"),
    4: ("APTO AL GRADO", "DECLARADO APTO", "DECLARAR APTO"),
    5: ("DICTAMEN DE TESIS", "DICTAMINANTE DE TESIS", "DICTAMINANTES DE TESIS"),
    6: ("FECHA Y HORA", "SUSTENTACION DE TESIS", "SUSTENTAR LA TESIS"),
    7: ("ELEVAR AL VRAC", "DIPLOMA DE GRADO", "OTORGAMIENTO DEL DIPLOMA", "OTORGAR EL GRADO ACADEMICO"),
}


def _primer_bloque_resolucion(texto: str) -> str:
    """Aísla la primera resolución cuando un Word contiene varias concatenadas."""
    texto = texto_normalizado(texto)
    inicios = [coincidencia.start() for coincidencia in re.finditer(r"ESCUELA DE POSGRADO RESOLUCION", texto)]
    return texto[: inicios[1]] if len(inicios) > 1 else texto


def _clasificar_paso(nombre: str, texto: str) -> int | None:
    """Pondera el objeto y RESUELVE; no usa la primera mención reglamentaria."""
    nombre = texto_normalizado(nombre)
    texto = _primer_bloque_resolucion(texto)
    inicio = texto[:1800]
    resolutiva = texto.split("RESUELVE", 1)[-1][:3500] if "RESUELVE" in texto else ""
    # Variantes cuyo nombre operativo describe mejor el objeto que las citas
    # reglamentarias repetidas dentro del documento.
    if "CAMBIO DE ASESOR" in nombre:
        return 1
    if "CAMBIO" in nombre and "DICTAM" in nombre:
        evidencia = f"{nombre} {inicio} {resolutiva}"
        return 2 if "PROYECTO" in evidencia else 5
    if "AMPLIACION" in nombre and any(variante in nombre for variante in ("PROYECTO", "PORYECTO")):
        return 3
    puntajes_nombre = Counter()
    puntajes = Counter()
    for paso, patrones in PATRONES_PASO.items():
        for patron in patrones:
            puntajes_nombre[paso] += nombre.count(patron)
            puntajes[paso] += 12 * nombre.count(patron)
            puntajes[paso] += 4 * inicio.count(patron)
            puntajes[paso] += 7 * resolutiva.count(patron)
            puntajes[paso] += texto.count(patron)
    # "Dictamen" aparece como antecedente de inscripción y apto; el verbo
    # resolutivo y el nombre del archivo deben prevalecer sobre esa referencia.
    if "INSCRIBIR" in resolutiva and "PROYECTO" in resolutiva:
        puntajes[3] += 18
    if "DECLARAR" in resolutiva and "APTO" in resolutiva:
        puntajes[4] += 18
    if "FIJAR" in resolutiva and "SUSTENTACION" in resolutiva:
        puntajes[6] += 18
    if "ELEVAR" in resolutiva and "VRAC" in resolutiva:
        puntajes[7] += 22
    # Los nombres operativos de Secretaría suelen expresar el objeto real y
    # son más fiables que antecedentes reglamentarios repetidos en el cuerpo.
    if puntajes_nombre:
        return puntajes_nombre.most_common(1)[0][0]
    if not puntajes:
        return None
    paso, puntaje = puntajes.most_common(1)[0]
    return paso if puntaje >= 7 else None


def _clasificar_variante(nombre: str, texto: str, paso: int) -> str:
    nombre = texto_normalizado(nombre)
    texto = _primer_bloque_resolucion(texto)
    resolutiva = texto.split("RESUELVE", 1)[-1][:3500] if "RESUELVE" in texto else ""
    evidencia = f"{nombre} {texto[:1800]} {resolutiva}"
    if "CAMBIO DE ASESOR" in nombre or "CAMBIO DE NOMBRAMIENTO DE ASESOR" in nombre:
        return "cambio_asesor"
    if "CAMBIO" in nombre and "DICTAM" in nombre:
        return "cambio_dictaminante"
    if "RECTIFIC" in nombre or "ERROR MATERIAL" in nombre:
        return "rectificacion"
    if "DEJAR SIN EFECTO" in nombre:
        return "dejar_sin_efecto"
    if "AMPLIACION" in nombre:
        return "ampliacion"
    if "RENUNCIA" in nombre:
        return "renuncia"
    if "RECTIFIC" in evidencia or "ERROR MATERIAL" in evidencia:
        return "rectificacion"
    if "DEJAR SIN EFECTO" in evidencia:
        return "dejar_sin_efecto"
    if "CAMBIO DE ASESOR" in evidencia or "CAMBIO DE NOMBRAMIENTO DE ASESOR" in evidencia:
        return "cambio_asesor"
    if "CAMBIO DE DICTAM" in evidencia:
        return "cambio_dictaminante"
    if "AMPLIACION" in evidencia:
        return "ampliacion"
    if "RENUNCIA" in evidencia:
        return "renuncia"
    if "CURSO DE ACTUALIZACION" in texto or re.search(r"\bC\.?A\.?I\.?\b", texto):
        return "cai"
    if paso == 7 and "CEPG-UAC" in texto:
        return "consejo_epg"
    return "regular"


def _puntaje_candidato(ruta: Path, texto: str, paso: int, variante: str) -> int:
    nombre = texto_normalizado(ruta.name)
    camino = texto_normalizado(str(ruta.parent))
    puntaje = 0
    puntaje += 35 if "PLANTILLA" in nombre else 0
    puntaje += 18 if "PLANTILLAS DE RESOLUCIONES" in camino else 0
    puntaje += 12 if "2026" in nombre else 0
    puntaje += 8 if re.search(r"(?:0000|XXXX|00XX)", nombre) else 0
    puntaje += 6 if 1400 <= len(texto) <= 9000 else 0
    puntaje += 5 * sum(marca in texto for marca in ("VISTO", "CONSIDERANDO", "RESUELVE"))
    puntaje += 6 if variante == "regular" else 0
    puntaje += 3 if NOMBRES_PASO[paso].split()[0].upper() in nombre else 0
    return puntaje


def analizar_archivo(ruta: Path, raiz: Path) -> dict:
    if ruta.name.startswith("~$"):
        return {"ruta": str(ruta), "estado": "temporal_word"}
    try:
        contenido = ruta.read_bytes()
        documento = Document(ruta)
        texto_crudo = _texto_documento(documento)
    except Exception as exc:
        return {"ruta": str(ruta), "estado": "ilegible", "error": str(exc)}

    texto = texto_normalizado(f"{ruta.name}\n{texto_crudo}")
    # Actas, oficios, informes y constancias viven en la carpeta, pero no son
    # modelos de resolución del circuito de siete pasos.
    es_resolucion = "RESOLUCION" in texto and "RESUELVE" in texto
    paso = _clasificar_paso(ruta.name, texto_crudo) if es_resolucion else None
    variante = _clasificar_variante(ruta.name, texto, paso) if paso else None
    return {
        "ruta": str(ruta),
        "ruta_relativa": str(ruta.relative_to(raiz)),
        "nombre": ruta.name,
        "estado": "candidata" if paso else "fuera_catalogo",
        "es_resolucion": es_resolucion,
        "id_paso": paso,
        "paso": NOMBRES_PASO.get(paso),
        "variante": variante,
        "es_cai": variante == "cai",
        "tiene_plantilla_en_nombre": "PLANTILLA" in texto_normalizado(ruta.name),
        "puntaje": _puntaje_candidato(ruta, texto, paso, variante) if paso else 0,
        "tamano": len(contenido),
        "caracteres_texto": len(texto_crudo),
        "sha256": hashlib.sha256(contenido).hexdigest(),
        "huella_formato": _huella_formato(documento),
        "preview": re.sub(r"\s+", " ", texto_crudo).strip()[:700],
    }


def construir_catalogo(raiz: Path, salida: Path) -> dict:
    archivos = sorted(ruta for ruta in raiz.rglob("*.docx") if ruta.is_file())
    resultados = [analizar_archivo(ruta, raiz) for ruta in archivos]
    candidatas = [item for item in resultados if item.get("estado") == "candidata"]

    grupos = defaultdict(list)
    for item in candidatas:
        grupos[(item["id_paso"], item["variante"])].append(item)

    canonicas = []
    carpeta_canonicas = salida / "canonicas"
    carpeta_canonicas.mkdir(parents=True, exist_ok=True)
    for (paso, variante), items in sorted(grupos.items()):
        formatos = Counter(item["huella_formato"] for item in items)
        for item in items:
            item["frecuencia_formato"] = formatos[item["huella_formato"]]
            item["puntaje_final"] = item["puntaje"] + min(item["frecuencia_formato"], 15)
        elegido = max(
            items,
            key=lambda item: (item["puntaje_final"], item["caracteres_texto"], item["nombre"]),
        )
        codigo = f"P{paso}_{variante.upper()}"
        destino = carpeta_canonicas / f"{codigo}.docx"
        shutil.copy2(elegido["ruta"], destino)
        canonicas.append(
            {
                "codigo": codigo,
                "id_paso": paso,
                "paso": NOMBRES_PASO[paso],
                "variante": variante,
                "archivo": str(destino),
                "origen": elegido["ruta"],
                "puntaje": elegido["puntaje_final"],
                "copias_clasificadas": len(items),
                "formatos_detectados": len(formatos),
                "preview": elegido["preview"],
            }
        )

    resumen = {
        "generado_en": datetime.now().isoformat(timespec="seconds"),
        "raiz_analizada": str(raiz),
        "archivos_docx": len(archivos),
        "resoluciones_clasificadas": len(candidatas),
        "fuera_catalogo": sum(item.get("estado") == "fuera_catalogo" for item in resultados),
        "ilegibles": sum(item.get("estado") == "ilegible" for item in resultados),
        "temporales_word": sum(item.get("estado") == "temporal_word" for item in resultados),
        "modelos_canonicos": canonicas,
        "cobertura_base": {
            str(paso): any(item["id_paso"] == paso and item["variante"] == "regular" for item in canonicas)
            for paso in NOMBRES_PASO
        },
    }
    salida.mkdir(parents=True, exist_ok=True)
    (salida / "catalogo.json").write_text(
        json.dumps(resumen, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (salida / "inventario.jsonl").write_text(
        "\n".join(json.dumps(item, ensure_ascii=False) for item in resultados) + "\n",
        encoding="utf-8",
    )
    lineas = [
        "# Catálogo de plantillas de resolución",
        "",
        f"Generado: {resumen['generado_en']}",
        f"Archivos DOCX revisados: {len(archivos)}",
        f"Resoluciones clasificadas: {len(candidatas)}",
        f"Modelos canónicos: {len(canonicas)}",
        "",
        "| Código | Paso | Variante | Copias | Formatos | Archivo origen |",
        "|---|---:|---|---:|---:|---|",
    ]
    lineas.extend(
        f"| {item['codigo']} | {item['id_paso']} | {item['variante']} | "
        f"{item['copias_clasificadas']} | {item['formatos_detectados']} | {Path(item['origen']).name} |"
        for item in canonicas
    )
    (salida / "REPORTE_CATALOGO.md").write_text("\n".join(lineas) + "\n", encoding="utf-8")
    return resumen


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--raiz", type=Path, default=ROOT_PREDETERMINADO)
    parser.add_argument("--salida", type=Path, default=SALIDA_PREDETERMINADA)
    args = parser.parse_args()
    if not args.raiz.exists():
        raise SystemExit(f"No existe la carpeta: {args.raiz}")
    resumen = construir_catalogo(args.raiz, args.salida)
    print(json.dumps(resumen, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
