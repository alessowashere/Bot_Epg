"""Borradores Word trazables a partir de modelos de resolución institucionales."""

from __future__ import annotations

import hashlib
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from resoluciones_pipeline import (
    detectar_alumno_codigo,
    detectar_programa,
    detectar_resolucion,
    detectar_titulo,
)


MESES = (
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
)
MARCAS_NO_MODELO = ("DEJAR SIN EFECTO", "RECTIFIC", "CAMBIO DE ", "RENUNCIA", "ERROR MATERIAL")


def es_modelo_utilizable(modelo) -> bool:
    """Evita usar rectificaciones o cambios como patrón de una resolución nueva."""
    # Solo se examina el encabezado del texto: un título de tesis puede contener
    # palabras como "cambio" sin que la resolución sea una rectificación.
    texto = " ".join(filter(None, [
        modelo.tipo_resolucion,
        modelo.archivo_normalizado,
        modelo.source_path,
        (modelo.texto_preview or "")[:1200],
    ]))
    texto = texto.upper()
    return not any(marca in texto for marca in MARCAS_NO_MODELO)


def fecha_literal(fecha: datetime) -> str:
    return f"Cusco, {fecha.day:02d} de {MESES[fecha.month - 1]} de {fecha.year}"


def serializar_modelo(modelo):
    return {
        "id_documento": modelo.id_documento,
        "numero_resolucion": modelo.resolucion_numero,
        "fecha_resolucion": modelo.fecha_resolucion.strftime("%Y-%m-%d") if modelo.fecha_resolucion else None,
        "tipo_resolucion": modelo.tipo_resolucion,
        "id_paso": modelo.id_paso_inferido,
        "estudiante_referencia": modelo.nombre_alumno,
        "codigo_referencia": modelo.codigo_alumno,
        "titulo_referencia": modelo.titulo_tesis,
        "archivo_referencia": modelo.archivo_normalizado or modelo.source_path,
        "texto_preview": modelo.texto_preview or "",
    }


def construir_borrador(tramite, modelo, numero_resolucion: str | None = None, fecha_resolucion: datetime | None = None):
    """Sustituye solo campos seguros; el resultado siempre requiere revisión humana."""
    expediente = tramite.expediente
    texto = (modelo.texto_preview or "").strip()
    if not texto:
        texto = (
            "ESCUELA DE POSGRADO\n"
            "RESOLUCIÓN N° [NÚMERO]\n"
            "[FECHA]\n\n"
            "VISTO: El expediente presentado por [ESTUDIANTE].\n\n"
            "CONSIDERANDO: [REVISAR FUNDAMENTO Y ANTECEDENTES].\n\n"
            "RESUELVE: [REDACTAR PARTE RESOLUTIVA SEGÚN EL MODELO APROBADO]."
        )

    numero = (numero_resolucion or tramite.numero_resolucion or "[NÚMERO DE RESOLUCIÓN]").strip()
    fecha = fecha_resolucion or tramite.fecha_resolucion
    titulo = expediente.titulo_tesis or "[TÍTULO DE TESIS POR COMPLETAR]"
    reemplazos = [
        (modelo.nombre_alumno, expediente.nombre_alumno),
        (modelo.codigo_alumno, expediente.codigo_alumno),
        (modelo.titulo_tesis, titulo),
    ]
    for origen, destino in reemplazos:
        if origen and destino:
            texto = re.sub(re.escape(origen), destino, texto, flags=re.IGNORECASE)

    # Solo la cabecera: referencias históricas dentro del considerando deben preservarse.
    texto = re.sub(r"(RESOLUCI[ÓO]N\s*N[°º.]?\s*)[^\n]{1,80}", rf"\g<1>{numero}", texto, count=1, flags=re.IGNORECASE)
    if fecha:
        texto = re.sub(r"Cusco\s*,\s*[^\n]{1,70}?20\d{2}", fecha_literal(fecha), texto, count=1, flags=re.IGNORECASE)

    # Algunos PDF se extraen como un único párrafo. Recuperar los hitos evita
    # entregar una vista previa o Word plano sin alterar el contenido jurídico.
    texto = re.sub(r"\s+(VISTO\s*[.:\-])", r"\n\n\1", texto, flags=re.IGNORECASE)
    texto = re.sub(r"\s+(CONSIDERANDO(?:S)?\s*[.:\-])", r"\n\n\1", texto, flags=re.IGNORECASE)
    texto = re.sub(r"\s+(RESUELVE\s*[.:\-])", r"\n\n\1", texto, flags=re.IGNORECASE)

    advertencias = [
        "Borrador generado desde un modelo histórico: revisar VISTO, CONSIDERANDO, RESUELVE, referencias y autoridades antes de remitir.",
    ]
    if not expediente.titulo_tesis:
        advertencias.append("El expediente no tiene título de tesis; el borrador conserva un marcador obligatorio.")
    if not expediente.codigo_alumno:
        advertencias.append("El expediente no tiene código de alumno; verificar la identificación antes de usar el borrador.")
    if len(modelo.texto_preview or "") < 1000:
        advertencias.append("El texto del modelo es parcial; completar el contenido institucional antes de remitir.")
    return {"contenido": texto, "advertencias": advertencias, "modelo": serializar_modelo(modelo)}


def crear_docx_borrador(contenido: str, tramite, modelo) -> bytes:
    documento = Document()
    seccion = documento.sections[0]
    seccion.top_margin = Pt(54)
    seccion.bottom_margin = Pt(54)
    encabezado = documento.add_paragraph()
    encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = encabezado.add_run("BORRADOR DE TRABAJO - REVISIÓN OBLIGATORIA")
    run.bold = True
    run.font.size = Pt(9)

    referencia = documento.add_paragraph()
    referencia.alignment = WD_ALIGN_PARAGRAPH.CENTER
    referencia.add_run(
        f"Modelo institucional: Resolución {modelo.resolucion_numero or 'sin número'} · "
        f"Paso {tramite.id_paso} · Trámite {tramite.uuid}"
    ).italic = True
    for bloque in re.split(r"\n\s*\n", contenido.strip()):
        parrafo = documento.add_paragraph()
        parrafo.paragraph_format.space_after = Pt(8)
        texto = bloque.strip()
        if texto.upper() in {"VISTO:", "CONSIDERANDO:", "RESUELVE:"}:
            parrafo.add_run(texto).bold = True
        else:
            parrafo.add_run(texto)
    salida = BytesIO()
    documento.save(salida)
    return salida.getvalue()


def hash_contenido(contenido: bytes) -> str:
    return hashlib.sha256(contenido).hexdigest()


def _parrafos_documento(documento):
    for parrafo in documento.paragraphs:
        yield parrafo
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                yield from celda.paragraphs
    for seccion in documento.sections:
        yield from seccion.header.paragraphs
        yield from seccion.footer.paragraphs


def texto_docx(ruta: Path) -> str:
    documento = Document(ruta)
    return "\n".join(parrafo.text for parrafo in _parrafos_documento(documento) if parrafo.text.strip())


def _reemplazar_en_parrafo(parrafo, patron, reemplazo, flags=re.IGNORECASE, limite=0):
    """Reemplaza texto incluso cuando Word lo partió en varios runs."""
    runs = list(parrafo.runs)
    texto = "".join(run.text for run in runs)
    coincidencias = list(re.finditer(patron, texto, flags=flags))
    if limite:
        coincidencias = coincidencias[:limite]
    for coincidencia in reversed(coincidencias):
        posiciones = []
        desplazamiento = 0
        for indice, run in enumerate(runs):
            posiciones.append((desplazamiento, desplazamiento + len(run.text), indice))
            desplazamiento += len(run.text)
        inicio_run = next((item for item in posiciones if item[0] <= coincidencia.start() < item[1]), None)
        fin_indice = max(coincidencia.end() - 1, coincidencia.start())
        fin_run = next((item for item in posiciones if item[0] <= fin_indice < item[1]), None)
        if not inicio_run or not fin_run:
            continue
        inicio_abs, _, indice_inicio = inicio_run
        inicio_fin_abs, _, indice_fin = fin_run
        prefijo = runs[indice_inicio].text[: coincidencia.start() - inicio_abs]
        sufijo = runs[indice_fin].text[coincidencia.end() - inicio_fin_abs :]
        runs[indice_inicio].text = prefijo + reemplazo + (sufijo if indice_inicio == indice_fin else "")
        for indice in range(indice_inicio + 1, indice_fin):
            runs[indice].text = ""
        if indice_fin != indice_inicio:
            runs[indice_fin].text = sufijo
    return len(coincidencias)


def referencias_plantilla_oficial(ruta: Path) -> dict:
    texto = texto_docx(ruta)
    nombre, codigo, criterio = detectar_alumno_codigo(texto, ruta.name)
    numero, anio = detectar_resolucion(texto, ruta.name)
    return {
        "nombre_alumno": nombre or None,
        "codigo_alumno": codigo or None,
        "criterio_alumno": criterio or None,
        "programa": detectar_programa(texto) or None,
        "titulo_tesis": detectar_titulo(texto) or None,
        "numero_resolucion": numero or None,
        "anio_resolucion": anio,
        "texto": texto,
    }


def crear_docx_desde_plantilla_oficial(ruta: Path, tramite, numero: str, fecha: datetime) -> tuple[bytes, dict]:
    """Conserva el DOCX institucional y sustituye únicamente campos identificados."""
    documento = Document(ruta)
    referencias = referencias_plantilla_oficial(ruta)
    expediente = tramite.expediente
    reemplazos = []
    for origen, destino, campo in (
        (referencias["nombre_alumno"], expediente.nombre_alumno, "estudiante"),
        (referencias["codigo_alumno"], expediente.codigo_alumno, "codigo"),
        (referencias["programa"], expediente.programa, "programa"),
        (referencias["titulo_tesis"], expediente.titulo_tesis, "titulo_tesis"),
    ):
        if origen and destino and origen.strip().upper() != destino.strip().upper():
            reemplazos.append((re.escape(origen.strip()), destino.strip(), campo, 0))

    # El número y fecha sólo se cambian en su primera aparición para preservar
    # las resoluciones históricas citadas en VISTO/CONSIDERANDO.
    reemplazos.extend(
        [
            (r"(?:0{2,4}|X{2,4}|\d{1,4})\s*[-/]\s*20\d{2}\s*(?:[-/]\s*(?:CEPG|EPG)[-\s/]*UAC)?", numero, "numero", 1),
            (r"Cusco\s*,\s*[^\n]{1,65}?20\d{2}", fecha_literal(fecha), "fecha", 1),
        ]
    )
    conteos = {campo: 0 for _, _, campo, _ in reemplazos}
    for patron, destino, campo, limite_total in reemplazos:
        restante = limite_total
        for parrafo in _parrafos_documento(documento):
            limite = restante if limite_total else 0
            cambios = _reemplazar_en_parrafo(parrafo, patron, destino, limite=limite)
            conteos[campo] += cambios
            if limite_total:
                restante -= cambios
                if restante <= 0:
                    break

    salida = BytesIO()
    documento.save(salida)
    advertencias = []
    for campo in ("estudiante", "codigo", "numero", "fecha"):
        if campo in conteos and not conteos[campo]:
            advertencias.append(f"No se pudo sustituir automáticamente el campo {campo}; revisarlo en Word.")
    # Algunas familias (rectificaciones, CAI, cambios de integrante, etc.) no
    # llevan título. Sólo advertimos cuando el modelo sí contenía ese campo.
    if not expediente.titulo_tesis and referencias["titulo_tesis"]:
        advertencias.append("El expediente no tiene título de tesis registrado.")
    return salida.getvalue(), {
        "referencias": {clave: valor for clave, valor in referencias.items() if clave != "texto"},
        "campos_reemplazados": conteos,
        "advertencias": advertencias,
    }
